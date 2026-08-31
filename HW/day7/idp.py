"""
Day 7 HW: IDP (Intelligent Document Processing)
把 data/ 底下的 1.pdf, 2.pdf, 3.pdf, 4.png, 5.docx 轉成純文字。

- 有文字層的 PDF（1.pdf, 2.pdf）：直接用 pypdf 抽取文字。
- 沒有文字層的掃描 PDF（3.pdf）：用 pypdfium2 把每一頁轉成圖片，
  再送進 Gemini 多模態視覺模型做 OCR。
- 圖片（4.png）：直接送進 Gemini 視覺模型做 OCR。
- docx（5.docx）：用 python-docx 抽取段落與表格文字。

抽取結果會快取到 output/idp_cache.json，避免重複呼叫 Gemini 視覺 API。
"""

import base64
import io
import json
import os
from pathlib import Path

import pypdf
import pypdfium2 as pdfium
from docx import Document
from dotenv import load_dotenv
from openai import OpenAI
from PIL import Image

load_dotenv(Path(__file__).parent / ".env")
load_dotenv(Path(__file__).parent.parent / ".env")

DATA_DIR = Path(__file__).parent / "data"
OUTPUT_DIR = Path(__file__).parent / "output"
CACHE_PATH = OUTPUT_DIR / "idp_cache.json"

GEMINI_BASE_URL = "https://generativelanguage.googleapis.com/v1beta/openai/"
GEMINI_MODEL = os.getenv("GEMINI_MODEL", "gemini-3.6-flash")

OCR_PROMPT = (
    "請將這張圖片中的所有文字逐字轉錄出來，保留原始語言（繁體中文），"
    "不要翻譯、不要摘要、不要加註解，也不要省略任何一行文字。"
)


def _vision_client() -> OpenAI:
    api_key = os.getenv("GEMINI_API_KEY")
    if not api_key:
        raise RuntimeError("缺少 GEMINI_API_KEY，請在 STEP/.env 中設定。")
    return OpenAI(base_url=GEMINI_BASE_URL, api_key=api_key, max_retries=8, timeout=180.0)


def _image_to_data_url(img: Image.Image) -> str:
    if img.mode != "RGB":
        img = img.convert("RGB")
    img.thumbnail((2000, 2000))
    buf = io.BytesIO()
    img.save(buf, format="JPEG", quality=90)
    b64 = base64.b64encode(buf.getvalue()).decode()
    return f"data:image/jpeg;base64,{b64}"


def _ocr_image(client: OpenAI, img: Image.Image) -> str:
    resp = client.chat.completions.create(
        model=GEMINI_MODEL,
        messages=[{
            "role": "user",
            "content": [
                {"type": "text", "text": OCR_PROMPT},
                {"type": "image_url", "image_url": {"url": _image_to_data_url(img)}},
            ],
        }],
        max_tokens=12000,
    )
    return resp.choices[0].message.content or ""


def _extract_pdf_text_layer(path: Path) -> list:
    """回傳每頁文字；若某頁沒有文字層，該頁回傳空字串。"""
    reader = pypdf.PdfReader(str(path))
    return [(page.extract_text() or "").strip() for page in reader.pages]


def _extract_pdf_via_ocr(client: OpenAI, path: Path) -> list:
    pdf = pdfium.PdfDocument(str(path))
    pages_text = []
    for i in range(len(pdf)):
        bitmap = pdf[i].render(scale=2.0)
        img = bitmap.to_pil()
        pages_text.append(_ocr_image(client, img).strip())
    return pages_text


def _extract_docx_text(path: Path) -> str:
    doc = Document(str(path))
    parts = [p.text for p in doc.paragraphs if p.text.strip()]
    for table in doc.tables:
        for row in table.rows:
            cells = [c.text for c in row.cells if c.text.strip()]
            if cells:
                parts.append(" | ".join(cells))
    return "\n".join(parts)


def extract_all(force: bool = False) -> dict:
    """
    處理 data/ 底下全部檔案，回傳 {doc_id: {"pages": [str, ...], "method": str}}。
    有快取的話直接讀快取（除非 force=True）。
    """
    if CACHE_PATH.exists() and not force:
        return json.loads(CACHE_PATH.read_text(encoding="utf-8"))

    client = _vision_client()
    result = {}

    # 1.pdf, 2.pdf：有文字層，直接抽取
    for name in ["1.pdf", "2.pdf"]:
        pages = _extract_pdf_text_layer(DATA_DIR / name)
        result[name] = {"pages": pages, "method": "pdf_text_layer"}

    # 3.pdf：掃描檔，逐頁用 Gemini 視覺 OCR
    pages = _extract_pdf_via_ocr(client, DATA_DIR / "3.pdf")
    result["3.pdf"] = {"pages": pages, "method": "gemini_vision_ocr"}

    # 4.png：圖片，直接 OCR
    img = Image.open(DATA_DIR / "4.png")
    text = _ocr_image(client, img).strip()
    result["4.png"] = {"pages": [text], "method": "gemini_vision_ocr"}

    # 5.docx：文字檔，直接抽取
    text = _extract_docx_text(DATA_DIR / "5.docx")
    result["5.docx"] = {"pages": [text], "method": "docx_text"}

    OUTPUT_DIR.mkdir(exist_ok=True)
    CACHE_PATH.write_text(json.dumps(result, ensure_ascii=False, indent=2), encoding="utf-8")
    return result


if __name__ == "__main__":
    data = extract_all(force=True)
    for doc_id, info in data.items():
        total_chars = sum(len(p) for p in info["pages"])
        print(f"{doc_id}: method={info['method']} pages={len(info['pages'])} chars={total_chars}")
